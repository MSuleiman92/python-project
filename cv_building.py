# Build an Application that would ask us a series of questions. 
# Upon completion when we answer the questions the application would build a CV for you. 

# First thing we install a package that would allow us work with document to do that we would type pip3 (package manager for python)
# Basically we can download existing packages that other developer have written and made available to the public 

# To install the package > pip3 install python-docx |  afterward > Successfully installed lxml-4.8.0 python-docx-0.8.11


# Now we need to use the library we just installed. To use a library, we use the command

from docx import Document  #from library docx we need
from docx.shared import Inches 
document = Document()

document.add_heading('Curriculum Vitae', 0)

document.add_picture(
    'profile.jpg', 
    width = Inches(2.0)
    )

# In Alternative we talk input from the User using  input()
name = input('What is your name ? ')
phone_number =input('What is your phone number ?')
email = input ('What is a email ?')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email )

# About yourself
document.add_heading('About me')
about_me = input('Tell me about yourself ?')
document.add_paragraph(about_me)

# Skill 
document.add_heading ('Professional Skills')
skill = input ('Enter Skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'


while True: 
    has_more_skills = input ('Do you have more skills Yes or No ?')
    if has_more_skills.lower() == 'yes':

        skill =input ('Enter Skill')
        p = document.add_paragraph(skill)
        p.style ='List Bullet'

    else: 
        break 


# Work experience 
document.add_heading ('Work Experience')
p = document.add_paragraph()

company = input ('Enter company')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run (company + ' ').bold =True
p.add_run(from_date + '-' + to_date + '\n').italic =True

experience_details = input ('Describe your experience at '+ company + ' ')
p.add_run (experience_details)

# More experiences 

while True: 
    has_more_experiences =input ('Do you have more experiences? Yes or No ?')
    if has_more_experiences == 'yes':

        p = document.add_paragraph()
        company = input ('Enter company')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run (company + ' ').bold =True
        p.add_run(from_date + '-' + to_date + '\n').italic =True

        experience_details = input ('Describe your experience at '+ company +' ')
        p.add_run (experience_details)     
    else:   
        break
 
# Footer
section = document.sections[0]
footer =section.footer
p = footer.paragraphs[0]
p.text ="CV generated using Python on VS code"

document.save('cv.docx')